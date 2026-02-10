# M:\meeting\app.py
import wave 
import time 
import asyncio
import json
import logging
import uuid
from datetime import datetime
from contextlib import asynccontextmanager
from dotenv import load_dotenv 
import os
import numpy as np
from fastapi import (
    FastAPI,
    UploadFile,
    Query,
    Request,
    WebSocket,
    HTTPException,
    Depends,
    status,
    Body
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.openapi.utils import get_openapi
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from starlette.websockets import WebSocketDisconnect, WebSocketState
import inspect 
from typing import Optional, List, Dict, Any, AsyncIterator 
from io import BytesIO
import docx 
from docx.shared import Inches
from starlette.responses import StreamingResponse, JSONResponse 
from urllib.parse import quote 

# Import Milvus related DataType for Schema definition
from pymilvus import connections # Import connections for unified connection management

load_dotenv()

# Import models and permission related classes
from models import TranscriptEntry, User, UserRole, Meeting

# Import backend coordinator and monitor manager
from backend.BackendCoordinator import BackendCoordinator
from backend.connection_manager import ConnectionManager # Still need to import ConnectionManager class
from backend.monitor_manager import MonitorManager

# Import service layer components
from core.knowledge_engine.kg_builder import KnowledgeGraphBuilder 
from config.settings import settings 
from services.mongodb_manager import MongoDBManager 
from services.milvus_service import MilvusManager 
from services.llm_service import LLMModel 
from services.embedding_service import BGEEmbeddingModel 
from services.neo4j_service import Neo4jService 
from services.permission_service import PermissionService 
from core.speech_to_text.stt_processor import SpeechToTextProcessor 
from services.voiceprint_service import VoiceprintService 
from services.summary_service import SummaryService 
from core.data_processing.minute_generator import MinuteGenerator 


# Configure root logger
logging.basicConfig(
    level=logging.INFO, # Can be set to DEBUG for more detailed logs
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Global system statistics
system_stats = {
    "total_connections": 0,
    "active_connections": 0,
    "audio_chunks_processed": 0,
    "meetings_active": 0,
    "stt_model_loaded": False,
    "speaker_model_loaded": False, 
    "summary_model_loaded": False, 
    "llm_model_loaded": False,     
    "bge_model_loaded": False      
}

# MongoDB global variable removed
# mongodb_db: Optional[MongoDBManager] = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """
    Application lifecycle event handler.
    Responsible for initializing services and database connections on application startup,
    and performing cleanup on application shutdown.
    """
    logger.info("Application starting (Lifespan)...")
    
    # Declare instances within lifespan to ensure they are accessible in the finally block for closing
    mongodb_manager_instance: Optional[MongoDBManager] = None
    voice_milvus_manager_instance: Optional[MilvusManager] = None 
    meeting_milvus_manager_instance: Optional[MilvusManager] = None 
    neo4j_service_instance: Optional[Neo4jService] = None 
    llm_model_instance: Optional[LLMModel] = None
    bge_embedding_model_instance: Optional[BGEEmbeddingModel] = None
    kg_builder_instance: Optional[KnowledgeGraphBuilder] = None 
    permission_service_instance: Optional[PermissionService] = None 
    stt_processor_instance: Optional[SpeechToTextProcessor] = None
    voiceprint_service_instance: Optional[VoiceprintService] = None 
    summary_service_instance: Optional[SummaryService] = None 



    try:
        # 0. Initialize MongoDBManager
        logger.info("Initializing MongoDBManager...")
        mongodb_manager_instance = MongoDBManager(
            host=settings.MONGO_HOST,
            port=settings.MONGO_PORT,
            db_name=settings.MONGO_DB_NAME
        )
        await mongodb_manager_instance.connect()
        app.state.mongodb_manager = mongodb_manager_instance
        logger.info("🎉 MongoDBManager initialized and connected.")

        # 1.1 Initialize Voiceprint MilvusManager
        logger.info(f"Initializing Voiceprint MilvusManager (Collection: {settings.MILVUS_VOICE_COLLECTION_NAME})...")
        voice_milvus_manager_instance = MilvusManager(
            config=settings, 
            collection_name=settings.MILVUS_VOICE_COLLECTION_NAME, 
            schema_fields=settings.MILVUS_VOICE_SCHEMA_FIELDS # Use properties from settings
        )
        # Do not force overwrite voiceprint collection to ensure Schema is correct and data persists
        await voice_milvus_manager_instance.connect(overwrite_collection=False) # <-- Modified to False
        app.state.voice_milvus_manager = voice_milvus_manager_instance
        logger.info("Voiceprint MilvusManager initialized and connected.")

        # 1.2 Initialize Meeting Text Embedding MilvusManager
        logger.info(f"Initializing Meeting Text Embedding MilvusManager (Collection: {settings.MILVUS_MEETING_COLLECTION_NAME})...")
        meeting_milvus_manager_instance = MilvusManager(
            config=settings, 
            collection_name=settings.MILVUS_MEETING_COLLECTION_NAME, 
            schema_fields=settings.MILVUS_MEETING_SCHEMA_FIELDS # Use properties from settings
        )
        # Do not force overwrite meeting text embedding collection to ensure Schema is correct
        await meeting_milvus_manager_instance.connect(overwrite_collection=False) 
        app.state.meeting_milvus_manager = meeting_milvus_manager_instance
        logger.info("Meeting Text Embedding MilvusManager initialized and connected.")


        # 2. Initialize Neo4jService
        logger.info("Initializing Neo4jService...")
        neo4j_service_instance = Neo4jService(
            settings.NEO4J_URI,
            settings.NEO4J_USER,
            settings.NEO4J_PASSWORD
        )
        await neo4j_service_instance.connect() 
        app.state.neo4j_service = neo4j_service_instance 
        logger.info("Neo4jService initialized and connected.")
        
        # 3. Initialize LLM Model
        logger.info("Initializing LLM Model...")
        llm_model_instance = LLMModel(settings_obj=settings) 
        await llm_model_instance.load_model() 
        app.state.llm_model = llm_model_instance
        system_stats["llm_model_loaded"] = llm_model_instance.is_model_loaded() 
        logger.info(f"LLM Model initialized. Load status: {system_stats['llm_model_loaded']}")

        # 4. Initialize BGE Embedding Model
        logger.info("Initializing BGE Embedding Model...")
        bge_embedding_model_instance = BGEEmbeddingModel(settings_obj=settings) 
        await bge_embedding_model_instance.load_model() 
        app.state.bge_model = bge_embedding_model_instance
        system_stats["bge_model_loaded"] = bge_embedding_model_instance.is_model_loaded() 
        logger.info(f"BGE Embedding Model initialized. Load status: {system_stats['bge_model_loaded']}")

        # 5. Instantiate KnowledgeGraphBuilder (requires LLM and Neo4jService here)
        logger.info("Instantiating KnowledgeGraphBuilder...")
        kg_builder_instance = KnowledgeGraphBuilder(
            neo4j_service=neo4j_service_instance, 
            llm_model=llm_model_instance,
            settings_obj=settings
        ) 
        app.state.kg_builder = kg_builder_instance 
        logger.info("KnowledgeGraphBuilder instance created.")

        # Initialize Neo4j Schema (constraints and indexes) before populating sample data
        await app.state.neo4j_service.initialize_schema() 
        logger.info("Neo4j Database Schema initialized.")

        # Perform knowledge graph health check and populate sample data if database is empty
        # Knowledge graph sample data initialization removed (Hardcoding fix)
        logger.info("Knowledge graph sample data initialization skipped (Clean startup).")

        # 6. Initialize PermissionService
        logger.info("Initializing PermissionService...")
        permission_service_instance = PermissionService() 
        app.state.permission_service = permission_service_instance
        logger.info("PermissionService initialized.")

        # 7. Initialize STTService (now SpeechToTextProcessor)
        logger.info("Initializing SpeechToTextProcessor...")
        stt_processor_instance = SpeechToTextProcessor(settings_obj=settings) 
        await stt_processor_instance.load_model() 
        app.state.stt_processor = stt_processor_instance
        system_stats["stt_model_loaded"] = stt_processor_instance.is_model_loaded()
        logger.info(f"SpeechToTextProcessor initialized. Load status: {system_stats['stt_model_loaded']}")

        # 8. Initialize VoiceprintService (passing Voiceprint MilvusManager and MongoDBManager)
        logger.info("Initializing VoiceprintService...")
        voiceprint_service_instance = VoiceprintService(
            settings_obj=settings, 
            voice_milvus_manager=voice_milvus_manager_instance, 
            mongodb_manager=mongodb_manager_instance 
        ) 
        await voiceprint_service_instance.load_model() 
        app.state.voiceprint_service = voiceprint_service_instance
        system_stats["speaker_model_loaded"] = voiceprint_service_instance.is_model_loaded()
        logger.info(f"VoiceprintService initialized. Load status: {system_stats['speaker_model_loaded']}")

        # 9. Initialize SummaryService
        logger.info("Initializing SummaryService...")
        summary_service_instance = SummaryService(settings_obj=settings)
        await summary_service_instance.load_model() 
        app.state.summary_service = summary_service_instance
        system_stats["summary_model_loaded"] = summary_service_instance.is_model_loaded()
        logger.info(f"SummaryService initialized. Load status: {system_stats['summary_model_loaded']}")

        # 10. Initialize ConnectionManager (before BackendCoordinator)
        app.state.connection_manager = ConnectionManager(system_stats_ref=system_stats)
        logger.info("ConnectionManager initialized.")

        # 11. Instantiate BackendCoordinator, passing in ConnectionManager
        logger.info("Instantiating BackendCoordinator...")
        app.state.backend_coordinator = BackendCoordinator(
            settings_obj=settings,
            mongodb_manager=mongodb_manager_instance, 
            voice_milvus_manager=voice_milvus_manager_instance, 
            meeting_milvus_manager=meeting_milvus_manager_instance, 
            llm_model=llm_model_instance,
            bge_model=bge_embedding_model_instance,
            neo4j_service=neo4j_service_instance, 
            permission_service=permission_service_instance, 
            stt_processor=stt_processor_instance, 
            voiceprint_service=voiceprint_service_instance, 
            summary_service=summary_service_instance,
            connection_manager=app.state.connection_manager 
        )
        logger.info("BackendCoordinator instance created.")

        # 12. Initialize MonitorManager and set it in BackendCoordinator
        app.state.monitor_manager = MonitorManager(system_stats_ref=system_stats) 
        app.state.backend_coordinator.set_monitor_manager(app.state.monitor_manager) 
        logger.info("MonitorManager initialized and set.")

        logger.info(f"System startup status: STT Model Loaded: {system_stats['stt_model_loaded']}, Speaker Model Loaded: {system_stats['speaker_model_loaded']}, Summary Model Loaded: {system_stats['summary_model_loaded']}, LLM Model Loaded: {system_stats['llm_model_loaded']}, BGE Model Loaded: {system_stats['bge_model_loaded']}")

        # Broadcast initial system status to all connected monitoring clients
        if hasattr(app.state, 'monitor_manager') and app.state.monitor_manager:
            await app.state.monitor_manager.broadcast({
                "type": "system_status_update",
                "data": system_stats,
                "timestamp": datetime.now().isoformat()
            })
        
    except Exception as e:
        logger.critical(f"Service initialization failed (Lifespan): {type(e).__name__}: {e}", exc_info=True)
        # Ensure that all initialized services are properly closed if startup fails
        if mongodb_manager_instance and hasattr(mongodb_manager_instance, 'close') and callable(mongodb_manager_instance.close):
            await mongodb_manager_instance.close()
        if voice_milvus_manager_instance and hasattr(voice_milvus_manager_instance, 'close') and callable(voice_milvus_manager_instance.close):
            await voice_milvus_manager_instance.close()
        if meeting_milvus_manager_instance and hasattr(meeting_milvus_manager_instance, 'close') and callable(meeting_milvus_manager_instance.close):
            await meeting_milvus_manager_instance.close()
        if neo4j_service_instance and hasattr(neo4j_service_instance, 'close') and callable(neo4j_service_instance.close):
            await neo4j_service_instance.close()
        if llm_model_instance and hasattr(llm_model_instance, 'close') and callable(llm_model_instance.close):
            await llm_model_instance.close()
        if bge_embedding_model_instance and hasattr(bge_embedding_model_instance, 'close') and callable(bge_embedding_model_instance.close):
            await bge_embedding_model_instance.close()
        if stt_processor_instance and hasattr(stt_processor_instance, 'close') and callable(stt_processor_instance.close):
            await stt_processor_instance.close()
        if voiceprint_service_instance and hasattr(voiceprint_service_instance, 'close') and callable(voiceprint_service_instance.close):
            await voiceprint_service_instance.close()
        if summary_service_instance and hasattr(summary_service_instance, 'close') and callable(summary_service_instance.close):
            await summary_service_instance.close()
        # ConnectionManager instance also needs to be closed
        if hasattr(app.state, 'connection_manager') and app.state.connection_manager:
            await app.state.connection_manager.close_all_connections() # <-- Modified
            logger.info("ConnectionManager cleaned up all connections (startup failed).")
        raise RuntimeError(f"Application startup failed: {type(e).__name__}: {e}") 

    yield 

    # --- Application Shutdown Logic ---
    logger.info("Application shutting down (Lifespan)...")
    try:
        # Close BackendCoordinator (it should handle closing its internal services)
        if hasattr(app.state, 'backend_coordinator') and app.state.backend_coordinator:
            close_method = getattr(app.state.backend_coordinator, 'close_services', None) 
            if close_method and callable(close_method):
                await close_method()
                logger.info("BackendCoordinator closed.")
            else:
                logger.warning("BackendCoordinator not initialized or no callable close_services method, skipping shutdown.")
        
        # Close all services created in lifespan (ensure they are closed after BackendCoordinator)
        if mongodb_manager_instance and hasattr(mongodb_manager_instance, 'close') and callable(mongodb_manager_instance.close):
            await mongodb_manager_instance.close()
            logger.info("MongoDBManager service closed.")
        
        # Milvus connections are disconnected here centrally, as MilvusManager instances share the same underlying connection
        if connections.has_connection(settings.MILVUS_ALIAS):
            try:
                connections.disconnect(alias=settings.MILVUS_ALIAS)
                logger.info("Milvus connection closed.")
            except Exception as e:
                logger.error(f"Error disconnecting from Milvus: {e}", exc_info=True)

        if neo4j_service_instance and hasattr(neo4j_service_instance, 'close') and callable(neo4j_service_instance.close):
            await neo4j_service_instance.close()
            logger.info("Neo4jService closed.")

        if llm_model_instance and hasattr(llm_model_instance, 'close') and callable(llm_model_instance.close):
            await llm_model_instance.close()
            logger.info("LLM model closed.")

        if bge_embedding_model_instance and hasattr(bge_embedding_model_instance, 'close') and callable(bge_embedding_model_instance.close):
            await bge_embedding_model_instance.close()
            logger.info("BGE embedding model closed.")

        if stt_processor_instance and hasattr(stt_processor_instance, 'close') and callable(stt_processor_instance.close):
            await stt_processor_instance.close()
            logger.info("SpeechToTextProcessor closed.")

        if summary_service_instance and hasattr(summary_service_instance, 'close') and callable(summary_service_instance.close):
            await summary_service_instance.close()
            logger.info("SummaryService closed.")

        logger.info("All services and database connections successfully closed (Lifespan).")
    except Exception as e:
        logger.error(f"Error closing services (Lifespan): {str(e)}", exc_info=True)

# Create FastAPI application instance
app = FastAPI(
    title="Meeting Intelligent Analysis System API",
    description="Provides meeting audio processing, knowledge graph querying, real-time voiceprint recognition, and other features",
    version="1.0.0",
    docs_url="/docs", 
    redoc_url="/redoc", 
    lifespan=lifespan 
)

# After FastAPI app instance creation, explicitly set log levels for specific modules to DEBUG
logging.getLogger('services').setLevel(logging.DEBUG)
logging.getLogger('backend').setLevel(logging.DEBUG)
logging.getLogger('core').setLevel(logging.DEBUG)
logger.debug("Log levels for 'services', 'backend', 'core' modules set to DEBUG.")

# Configure CORS (Cross-Origin Resource Sharing) middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  
    allow_credentials=True, 
    allow_methods=["*"], 
    allow_headers=["*"], 
)

# Mount static files (CSS, JS, images)
app.mount("/static", StaticFiles(directory="static"), name="static")
# Configure Jinja2 template engine to render HTML pages
templates = Jinja2Templates(directory="templates")

# --- Pydantic Data Models for API Request/Response Validation ---
class TestRequest(BaseModel):
    text: str
    max_length: int = 150
    min_length: int = 30
    num_beams: int = 4

class TestResponse(BaseModel):
    status: str
    summary: str
    input_length: int

    summary_length: int
    device: str

class LLMTestRequest(BaseModel):
    prompt: str
    max_new_tokens: int = 200
    temperature: float = 0.7

class LLMTestResponse(BaseModel):
    status: str
    response: str
    input_length: int
    output_length: int

# Custom OpenAPI document generation function
def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema

    openapi_schema = get_openapi(
        title="Meeting Intelligent Analysis System API",
        description="Provides meeting audio processing, knowledge graph querying, real-time voiceprint recognition, and other features",
        version="1.0.0",
        routes=app.routes,
    )

    # Add contact information to the OpenAPI document
    openapi_schema["info"]["contact"] = {
        "name": "Technical Support",
        "email": "459880255@qq.com"
    }
    app.openapi_schema = openapi_schema
    return app.openapi_schema

app.openapi = custom_openapi 

import random

# --- HTTP Route Handlers ---

@app.get("/")
async def home(request: Request):
    """
    Root route for serving the main meeting interface HTML page.
    """
    return templates.TemplateResponse("meeting_ui.html", {"request": request})

@app.get("/monitor")
async def get_monitor(request: Request):
    """
    Serves the monitoring dashboard HTML page.
    """
    return templates.TemplateResponse("monitor.html", {"request": request})

@app.get("/get_voice_login_text")
async def get_voice_login_text():
    """
    Provides a random text string for voice login/registration.
    """
    sentences = [
        "今天天气真不错，我们出去散散步吧。",
        "这个项目的截止日期是下周五。",
        "请确认您已收到会议纪要。",
        "创新是推动我们前进的核心动力。",
        "客户的满意是我们最大的追求。"
    ]
    return {"text": random.choice(sentences)}

@app.post("/voice_login")
async def voice_login_endpoint(file: UploadFile):
    """
    Handles voice login and registration.
    It identifies the speaker from the audio. If the speaker is recognized,
    it logs them in. If not, it registers them as a new user.
    """
    backend_coordinator = app.state.backend_coordinator
    if not backend_coordinator:
        raise HTTPException(status_code=503, detail="Backend services are not initialized.")

    voiceprint_service = backend_coordinator.voiceprint_service
    mongodb_manager = backend_coordinator.mongodb_manager

    if not voiceprint_service or not mongodb_manager:
        raise HTTPException(status_code=503, detail="Voice or database services are not available.")

    try:
        audio_content = await file.read()
        if not audio_content:
            raise HTTPException(status_code=400, detail="Audio file is empty.")

        audio_data_np = np.frombuffer(audio_content, dtype=np.int16).astype(np.float32) / 32768.0
        if np.isnan(audio_data_np).any() or np.isinf(audio_data_np).any():
            raise HTTPException(status_code=400, detail="Invalid audio data (NaN/Inf).")

        identification_result = await voiceprint_service.identify_speaker(audio_data_np, settings.VOICE_SAMPLE_RATE)

        if identification_result and identification_result.get("is_known"):
            user_id = identification_result.get("user_id")
            logger.info(f"Voice login successful for user: {user_id}")
            user_profile = await mongodb_manager.get_user(user_id)
            if not user_profile:
                 raise HTTPException(status_code=404, detail=f"User {user_id} found in voiceprint DB but not in user DB.")
            return {
                "message": "Voice login successful!",
                "user": user_profile
            }
        else:
            logger.info("Unknown speaker. Proceeding with new user registration.")
            new_user_id = str(uuid.uuid4())
            new_username = f"用户_{new_user_id[:6]}"
            new_role = UserRole.REGISTERED_USER.value
            await voiceprint_service.register_voice(
                audio_data=audio_data_np,
                sample_rate=settings.VOICE_SAMPLE_RATE,
                user_id=new_user_id,
                username=new_username,
                role=new_role
            )
            new_user_profile = await mongodb_manager.get_user(new_user_id)
            return {
                "message": "New user registered successfully!",
                "user": new_user_profile
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Voice login process failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during voice login: {e}")

@app.post("/test/bart/summarize")
async def test_bart_summarize(request: TestRequest) -> TestResponse:
    """
    Tests the functionality of the BART summary generation model.
    Receives text input and returns the summary result.
    """
    try:
        if not hasattr(app.state, 'backend_coordinator') or \
           app.state.backend_coordinator is None or \
           not hasattr(app.state.backend_coordinator, 'meeting_assistant') or \
           app.state.backend_coordinator.meeting_assistant is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Summary service not initialized, please try again later."
            )
        if not hasattr(app.state, 'summary_service') or \
           app.state.summary_service is None or \
           not app.state.summary_service.is_model_loaded():
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="Summary model not loaded, please try again later."
            )
        summary = await app.state.backend_coordinator.meeting_assistant.generate_summary(
            request.text,
            max_length=request.max_length,
            min_length=request.min_length,
            num_beams=request.num_beams
        )
        if not summary or len(summary) < 10 or "..." in summary:
            logger.warning(f"Summary generation incomplete or empty, input length: {len(request.text)}, summary: {summary}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Summary generation incomplete or empty, please adjust parameters or check model status."
            )
        summary_device = "N/A"
        if hasattr(app.state, 'summary_service') and \
           app.state.summary_service:
            summary_device = str(app.state.summary_service.device)
        return TestResponse(
            status="success",
            summary=summary,
            device=summary_device,
            input_length=len(request.text),
            summary_length=len(summary)
        )
    except HTTPException: 
        raise
    except Exception as e:
        logger.error(f"BART summary generation failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Processing failed: {str(e)}"
        )

@app.post("/test/llm/generate")
async def test_llm_generate(request: LLMTestRequest) -> LLMTestResponse:
    """
    Tests the functionality of the LLM model.
    Receives prompt text and returns the generated result.
    """
    try:
        if not hasattr(app.state, 'llm_model') or \
           app.state.llm_model is None or \
           not app.state.llm_model.is_model_loaded():
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="LLM model not loaded, please try again later."
            )
        full_response_chunks = []
        async for chunk in app.state.llm_model.generate_text(
            request.prompt,
            max_new_tokens=request.max_new_tokens,
            temperature=request.temperature,
            stream=False
        ):
            full_response_chunks.append(chunk)
        response_text = "".join(full_response_chunks)
        return LLMTestResponse(
            status="success",
            response=response_text,
            input_length=len(request.prompt),
            output_length=len(response_text)
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"LLM generation failed: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Processing failed: {str(e)}"
        )

@app.get("/system/status")
async def system_status_endpoint():
    """
    Provides real-time system status information, including active connections, processed audio chunks, model load status, etc.
    """
    global system_stats
    if hasattr(app.state, 'connection_manager'):
        system_stats["active_connections"] = len(app.state.connection_manager.active_connections)
    if hasattr(app.state, 'backend_coordinator') and app.state.backend_coordinator:
        vad_service = getattr(app.state.backend_coordinator, 'vad_service', None)
        if vad_service and hasattr(vad_service, 'is_model_loaded') and callable(vad_service.is_model_loaded):
            system_stats["vad_model_loaded"] = vad_service.is_model_loaded()
        else:
            system_stats["vad_model_loaded"] = False
        stt_processor = getattr(app.state.backend_coordinator, 'stt_processor', None)
        if stt_processor and hasattr(stt_processor, 'is_model_loaded') and callable(stt_processor.is_model_loaded):
            system_stats["stt_model_loaded"] = stt_processor.is_model_loaded()
        else:
            system_stats["stt_model_loaded"] = False
        speaker_service = getattr(app.state.backend_coordinator, 'voiceprint_service', None)
        system_stats["speaker_model_loaded"] = True if speaker_service and speaker_service.is_model_loaded() else False
        summary_service = getattr(app.state.backend_coordinator, 'summary_service', None)
        if summary_service and hasattr(summary_service, 'is_model_loaded') and callable(summary_service.is_model_loaded):
            system_stats["summary_model_loaded"] = summary_service.is_model_loaded()
        else:
            system_stats["summary_model_loaded"] = False
        llm_model = getattr(app.state, 'llm_model', None)
        if llm_model and hasattr(llm_model, 'is_model_loaded') and callable(llm_model.is_model_loaded):
            system_stats["llm_model_loaded"] = llm_model.is_model_loaded()
        else:
            system_stats["llm_model_loaded"] = False
        bge_model = getattr(app.state, 'bge_model', None)
        if bge_model and hasattr(bge_model, 'is_model_loaded') and callable(bge_model.is_model_loaded):
            system_stats["bge_model_loaded"] = bge_model.is_model_loaded()
        else:
            system_stats["bge_model_loaded"] = False
    return {
        "status": "running",
        "meetings_active": system_stats["meetings_active"],
        "connections_active": system_stats["active_connections"],
        "audio_processed": system_stats["audio_chunks_processed"],
        "vad_model_loaded": system_stats["vad_model_loaded"],
        "stt_model_loaded": system_stats["stt_model_loaded"],
        "speaker_model_loaded": system_stats["speaker_model_loaded"],
        "summary_model_loaded": system_stats["summary_model_loaded"],
        "llm_model_loaded": system_stats["llm_model_loaded"],
        "bge_model_loaded": system_stats["bge_model_loaded"],
        "timestamp": datetime.now().isoformat(),
        "total_connections": system_stats["total_connections"]
    }

@app.get("/test/bart/status")
async def check_bart_status():
    """
    Checks if the BART summary model is loaded and ready.
    """
    if hasattr(app.state, 'backend_coordinator') and \
       app.state.backend_coordinator and \
       hasattr(app.state.backend_coordinator, 'summary_service') and \
       app.state.backend_coordinator.summary_service and \
       callable(getattr(app.state.backend_coordinator.summary_service, 'is_model_loaded', None)):
        is_loaded = app.state.backend_coordinator.summary_service.is_model_loaded()
        device_info = str(app.state.backend_coordinator.summary_service.device)
        return {
            "status": "running" if is_loaded else "initializing",
            "device": device_info,
            "model": settings.SUMMARY_MODEL_HUB_NAME, 
            "ready": is_loaded
        }
    return {
        "status": "initializing",
        "device": "N/A",
        "model": settings.SUMMARY_MODEL_HUB_NAME,
        "ready": False
    }

@app.post("/upload-meeting")
async def upload_meeting(audio: UploadFile, transcript: str):
    """
    接收完整的会议录音文件和预转录的文本。
    用于对整个会议进行离线处理，例如生成摘要、识别行动项等。
    """
    try:
        if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="后端服务未初始化，请稍后重试。"
            )
        audio_content = await audio.read()
        if not audio_content:
            logger.error("上传的音频文件为空。")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传的音频文件为空。"
            )
        try:
            audio_data_np = np.frombuffer(audio_content, dtype=np.int16).astype(np.float32) / 32768.0
        except ValueError as e:
            logger.error(f"无法将音频字节转换为 NumPy 数组: {e}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"音频文件格式不兼容或已损坏: {e}"
            )
        if np.isnan(audio_data_np).any() or np.isinf(audio_data_np).any():
            logger.error("转换后的音频数据包含 NaN 或 Inf 值，无法处理。")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传的音频文件包含无效数据 (NaN/Inf)。请检查文件是否损坏或格式不正确。"
            )
        result = await app.state.backend_coordinator.process_meeting(audio_data_np, settings.VOICE_SAMPLE_RATE, transcript)
        logger.info(f"会议音频处理完成。")
        return {
            "summary": result.get('summary', '未能生成摘要'),
            "actions": result.get('actions', []),
            "decision": result.get('decision', '没有重要决策'),
            "identified_speakers": result.get('identified_speakers', {})
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"上传会议录音失败: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"处理会议录音失败: {str(e)}"
        )

@app.get("/kg-health")
async def knowledge_graph_health():
    """
    检查知识图谱数据库的健康状态。
    """
    try:
        if not hasattr(app.state, 'backend_coordinator') or \
           app.state.backend_coordinator is None or \
           not hasattr(app.state.backend_coordinator, 'kg_builder') or \
           app.state.backend_coordinator.kg_builder is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="知识图谱服务未初始化，请稍后重试。"
            )
        is_empty = await app.state.backend_coordinator.kg_builder.is_database_empty()
        return {
            "status": "健康",
            "empty": is_empty,
            "schema_initialized": not is_empty
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"知识图谱健康检查失败: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"知识图谱健康检查失败: {str(e)}"
        )

@app.get("/query-knowledge")
async def query_knowledge(entity: str, depth: int = Query(default=1, ge=1, le=5)):
    """
    查询知识图谱中与给定实体相关的信息。
    """
    try:
        if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="后端服务未初始化，请稍后重试。"
            )
        return await app.state.backend_coordinator.query_knowledge_graph(entity, depth)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"知识图谱查询失败 (实体: '{entity}', 深度: {depth}): {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"知识图谱查询失败: {str(e)}"
        )

@app.post("/ask_llm_stream/")
async def ask_llm_stream(request: Dict[str, Any] = Body(..., example={"question": "什么是大语言模型？"})):
    """
    用于流式问答的 API 接口。
    接收一个包含 'question' 字段的 JSON 请求体。
    返回一个文本流。
    """
    question = request.get("question")
    if not question:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="请求体中必须提供 'question' 字段。"
        )
    async def generate_and_stream():
        try:
            if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
                logger.error("BackendCoordinator 未初始化，无法执行流式问答。")
                yield "服务器内部服务未准备就绪，无法回答您的问题。"
                return
            async for chunk in app.state.backend_coordinator.get_answer_from_llm(question):
                yield chunk
        except Exception as e:
            logger.error(f"流式问答生成过程中发生错误: {e}", exc_info=True)
            yield f"服务器内部错误: {str(e)}"
    return StreamingResponse(generate_and_stream(), media_type="text/plain")

@app.post("/register-voice")
async def register_voice_endpoint(
    audio: UploadFile,
    user_name: str = Query(..., description="要注册用户的友好名称"),
    role: UserRole = Query(..., description="用户的角色，例如 GUEST, MEMBER, ADMIN")
):
    """
    通过上传音频文件来注册用户的声纹。
    """
    try:
        # 检查 BackendCoordinator 是否已初始化
        if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="后端服务未初始化，请稍后重试。"
            )

        # 读取音频内容
        audio_content = await audio.read()
        if not audio_content:
            logger.error("上传的音频文件为空。")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传的音频文件为空。"
            )

        # 尝试将音频字节转换为 numpy 数组 (float32)
        # 假设原始音频是 int16 PCM，这是一种常见的音频格式
        try:
            audio_data_np = np.frombuffer(audio_content, dtype=np.int16).astype(np.float32) / 32768.0
        except ValueError as e:
            logger.error(f"无法将音频字节转换为 NumPy 数组: {e}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"音频文件格式不兼容或已损坏: {e}"
            )

        # 检查 audio_data_np 是否包含 NaN 或 Inf
        if np.isnan(audio_data_np).any() or np.isinf(audio_data_np).any():
            logger.error("转换后的音频数据包含 NaN 或 Inf 值，无法进行声纹注册。")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="上传的音频文件包含无效数据 (NaN/Inf)。请检查文件是否损坏或格式不正确。"
            )

        # 通过 BackendCoordinator 注册声纹
        result = await app.state.backend_coordinator.register_voice(audio_data_np, settings.VOICE_SAMPLE_RATE, user_name, role.value)

        if result.get("status") == "registered":
            logger.info(f"用户 '{user_name}' (角色: {role.value}) 声纹注册成功。原始 ID: {result.get('user_id')}")
            return {
                "message": "声纹注册成功",
                "user_id": result.get("user_id"),
                "user_name": user_name,
                "role": role.value
            }
        else:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=result.get("message", "声纹注册失败")
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"声纹注册失败: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"声纹注册处理失败: {str(e)}"
        )


# --- WebSocket Endpoint ---

@app.websocket("/ws/meeting/{meeting_id}/{client_id}")
async def meeting_websocket(websocket: WebSocket, meeting_id: str, client_id: str):
    """
    实时会议WebSocket连接。
    接收前端的音频块和控制消息，进行实时声纹识别和语音转文字，并将结果广播给客户端。
    """
    logger.info(f"正在尝试建立WebSocket连接 (客户端ID: {client_id}, 会议ID: {meeting_id})")

    # 确保BackendCoordinator已初始化
    if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
        logger.error("BackendCoordinator未初始化，无法建立WebSocket连接。")
        if websocket.client_state in (WebSocketState.CONNECTING, WebSocketState.CONNECTED):
            await websocket.close(code=status.WS_1011_INTERNAL_ERROR, reason="服务器内部服务未就绪。")
        return

    # 从BackendCoordinator获取实时音频处理器实例
    meeting_assistant = app.state.backend_coordinator.meeting_assistant
    if meeting_assistant is None:
        logger.error("SmartMeetingAssistant (会议助手) 未初始化，无法建立WebSocket连接。")
        if websocket.client_state in (WebSocketState.CONNECTING, WebSocketState.CONNECTED):
            await websocket.close(code=status.WS_1011_INTERNAL_ERROR, reason="服务器实时音频处理服务未就绪。")
        return

    # 用于在连接期间累积完整的音频数据
    full_audio_bytes = b''

    try:
        # 与ConnectionManager建立连接
        await app.state.connection_manager.connect(websocket, client_id, meeting_id)
        logger.info(f"客户端 {client_id} (会议ID: {meeting_id}) WebSocket连接已接受。")

        # 向客户端发送初始化消息以确认连接并提供ID
        user_id_from_backend = client_id
        user_role_from_backend = UserRole.GUEST.value
        username_from_backend = f"用户_{client_id[:6]}"

        mongodb_db = getattr(app.state, 'mongodb_manager', None)
        if mongodb_db:
            try:
                # 尝试从MongoDB获取用户配置
                user_profile = await mongodb_db.get_user(client_id)
                if user_profile:
                    user_id_from_backend = user_profile.user_id
                    user_role_from_backend = user_profile.role.value
                    username_from_backend = user_profile.username
                else:
                    # 如果客户端ID在数据库中不存在，则以访客身份添加
                    new_user = User(user_id=client_id, username=username_from_backend, role=UserRole.GUEST, last_active=datetime.utcnow())
                    await mongodb_db.add_or_update_user(new_user)
                    logger.info(f"MongoDB中未找到客户端ID {client_id} 的用户配置，已创建为默认访客角色。")
            except Exception as e:
                logger.warning(f"获取WebSocket用户信息失败 (MongoDB): {e}。将使用客户端ID作为用户ID。", exc_info=True)
        else:
            logger.warning("MongoDB数据库未初始化，无法获取WebSocket用户信息。")


        await app.state.connection_manager.send_personal_message({
            "type": "meeting_init_response",
            "status": "success",
            "meetingId": meeting_id,
            "clientId": client_id,
            "userId": user_id_from_backend,
            "role": user_role_from_backend,
            "username": username_from_backend,
            "timestamp": datetime.now().isoformat()
        }, client_id)
        logger.info(f"已向客户端 {client_id} 发送 meeting_init_response。")

        # 接收客户端消息的主循环
        while True:
            message = await websocket.receive()
            logger.debug(f"从客户端 {client_id} 收到原始 WebSocket 消息: {message}")

            # 首先检查消息类型是否为websocket.receive事件
            if message.get("type") == "websocket.receive":
                if "bytes" in message:
                    logger.debug(f"从客户端 {client_id} 收到二进制音频数据。")
                    audio_bytes = message["bytes"]
                    if not audio_bytes:
                        logger.warning(f"从客户端 {client_id} 接收到空音频字节。")
                        continue

                    # 将接收到的音频块追加到完整音频字节流中
                    full_audio_bytes += audio_bytes

                    if 'system_stats' in globals() and isinstance(system_stats, dict):
                        system_stats["audio_chunks_processed"] += 1

                    # 尝试将音频字节转换为numpy数组 (float32)
                    try:
                        audio_data_np = np.frombuffer(audio_bytes, dtype=np.int16).astype(np.float32) / 32768.0
                    except ValueError as e:
                        logger.error(f"将WebSocket音频字节转换为NumPy数组失败: {e}", exc_info=True)
                        await app.state.connection_manager.send_personal_message({
                            "type": "error",
                            "message": f"实时音频数据格式不兼容或已损坏: {e}",
                            "timestamp": datetime.now().isoformat()
                        }, client_id)
                        continue

                    # 检查实时音频数据是否包含NaN或Inf
                    if np.isnan(audio_data_np).any() or np.isinf(audio_data_np).any():
                        logger.error(f"实时音频数据包含NaN或Inf值，跳过处理 (客户端: {client_id}, 会议: {meeting_id})。")
                        await app.state.connection_manager.send_personal_message({
                            "type": "error",
                            "message": "实时音频数据无效，请检查麦克风或音频源。",
                            "timestamp": datetime.now().isoformat()
                        }, client_id)
                        continue

                    # 调用SmartMeetingAssistant的process_real_time_audio方法
                    await meeting_assistant.process_real_time_audio(audio_data_np, settings.VOICE_SAMPLE_RATE, client_id, meeting_id)

                elif "text" in message:
                    text_data = message["text"]
                    try:
                        json_data = json.loads(text_data)
                        if isinstance(json_data, dict):
                            logger.info(f"从客户端 {client_id} 接收到文本消息: {json_data.get('type', '未知类型')}")

                            if json_data.get("type") == "start_recording":
                                logger.info(f"客户端 {client_id} 发送 'start_recording' 命令。")
                            elif json_data.get("type") == "stop_recording":
                                logger.info(f"客户端 {client_id} 发送 'stop_recording' 命令。")
                            elif json_data.get("type") == "client_ready":
                                logger.info(f"客户端 {client_id} 已准备好接收数据。")
                            elif json_data.get("type") == "update_user_role":
                                user_id_to_update = json_data.get("user_id")
                                new_role_str = json_data.get("new_role")
                                if user_id_to_update and new_role_str:
                                    try:
                                        new_role = UserRole(new_role_str.upper())
                                        logger.info(f"WS: 收到将用户 {user_id_to_update} 角色更新为 {new_role.value} 的请求。")
                                        update_result = await app.state.backend_coordinator.update_user_role_from_ws(user_id_to_update, new_role.value) # 传递字符串值
                                        await app.state.connection_manager.send_personal_message(update_result, client_id)
                                    except ValueError:
                                        logger.warning(f"收到的角色值无效: {new_role_str}")
                                        await app.state.connection_manager.send_personal_message({
                                            "type": "error",
                                            "message": f"角色值无效: {new_role_str}",
                                            "timestamp": datetime.now().isoformat()
                                        }, client_id)
                                else:
                                    logger.warning(f"收到的 update_user_role 消息无效: {json_data}")
                            else:
                                logger.warning(f"从客户端 {client_id} 收到未知JSON消息类型: {json_data.get('type', '无类型字段')}")

                        else:
                            logger.warning(f"从客户端 {client_id} 收到非字典JSON消息: {text_data}")

                    except json.JSONDecodeError:
                        logger.warning(f"从客户端 {client_id} 收到非JSON文本消息: {text_data}")
                    except Exception as e:
                        logger.error(f"处理来自客户端 {client_id} 的文本消息时出错: {e}", exc_info=True)
                else:
                    logger.warning(f"从客户端 {client_id} 收到空的 'websocket.receive' 事件，既无字节也无文本。键: {message.keys()}")
            elif message.get("type") == "websocket.disconnect":
                logger.info(f"从客户端 {client_id} 收到 'websocket.disconnect' 消息，正在关闭连接。")
                break
            else:
                logger.warning(f"从客户端 {client_id} 收到非 'websocket.receive' 事件: {message.get('type', '未知')}。键: {message.keys()}")

    except WebSocketDisconnect:
        logger.info(f"客户端 {client_id} (会议ID: {meeting_id}) 断开连接。")
    except RuntimeError as e:
        if "disconnect message has been received" in str(e):
            logger.warning(f"客户端 {client_id} 的WebSocket连接已断开，但意外地再次尝试接收消息。")
        else:
            logger.error(f"WebSocket运行时错误 (客户端ID: {client_id}, 会议ID: {meeting_id}): {str(e)}", exc_info=True)
        try:
            await app.state.connection_manager.send_personal_message({
                "type": "error",
                "message": f"服务器错误: {str(e)}",
                "timestamp": datetime.now().isoformat()
            }, client_id)
        except Exception as send_e:
            logger.warning(f"向 {client_id} 发送错误消息失败: {send_e}")
    finally:
        
        try:
            if hasattr(app.state, 'connection_manager') and app.state.connection_manager is not None:
                if websocket.client_state == WebSocketState.CONNECTED:
                    # 确保只在连接仍处于“已连接”状态时才尝试断开连接
                    await app.state.connection_manager.disconnect(client_id, meeting_id, expected_websocket=websocket)
        except Exception as cleanup_e:
            logger.error(f"为客户端 {client_id} (会议ID: {meeting_id}) 清理资源时出错: {cleanup_e}", exc_info=True)
        logger.info(f"客户端 {client_id} 的资源清理已完成, 会议ID: {meeting_id}")

# --- Get All Users API ---
@app.get("/get-all-users", response_model=List[User])
async def get_all_users_endpoint():
    """
    获取所有已注册的用户及其角色。
    """
    mongodb_db = getattr(app.state, 'mongodb_manager', None)
    if not mongodb_db:
        logger.error("❌ MongoDB 数据库未初始化。无法获取用户列表。")
        raise HTTPException(status_code=500, detail="MongoDB 数据库未初始化。")

    try:
        users_data = await mongodb_db.get_all_users()
        return users_data
    except Exception as e:
        logger.error(f"获取所有用户失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取用户失败: {e}")

# --- 更新用户角色 API ---
@app.post("/update-role")
async def update_user_role_endpoint(
    voiceprint_id: str = Query(..., description="要更新的用户的声纹 ID"),
    new_role: UserRole = Query(..., description="用户的新角色，例如 GUEST、MEMBER、ADMIN"),
    new_name: str = Query(None, description="可选的用户新名称")
):
    """
    更新用户的角色和/或名称。
    需要当前用户拥有 `EDIT_ROLES` 权限。

    参数:
    - voiceprint_id: 用户的唯一标识符
    - new_role: 要分配的新角色
    - new_name: 可选的新名称（仅在提供时更新）

    返回:
    - 包含操作状态的 JSON 响应
    """
    mongodb_db = getattr(app.state, 'mongodb_manager', None)
    if not mongodb_db:
        raise HTTPException(status_code=500, detail="MongoDB 数据库未初始化。")

    try:
        # 验证用户存在
        existing_user = await mongodb_db.get_user(voiceprint_id)
        if not existing_user:
            raise HTTPException(
                status_code=404,
                detail=f"未找到 ID 为 '{voiceprint_id}' 的用户。"
            )

        # 执行更新操作
        update_success = await mongodb_db.update_user_role(
            voiceprint_id,
            new_role,
            new_name  # 传递新名称参数
        )

        if not update_success:
            raise HTTPException(
                status_code=500,
                detail="更新操作失败（没有匹配的文档）"
            )

        # 构造响应消息
        response_msg = f"用户 '{voiceprint_id}' 的角色已更新为 '{new_role.value}'"
        if new_name:
            response_msg += f" 且名称已更新为 '{new_name}'"

        logger.info(response_msg)
        return {
            "message": response_msg,
            "status": "success",
            "data": {
                "voiceprint_id": voiceprint_id,
                "new_role": new_role.value,
                "new_name": new_name if new_name else None
            }
        }

    except HTTPException as e:
        # 已知异常直接抛出
        raise e
    except Exception as e:
        logger.error(
            f"更新用户 '{voiceprint_id}' 失败: {str(e)}",
            exc_info=True
        )
        raise HTTPException(
            status_code=500,
            detail=f"更新用户失败: {str(e)}"
        )

# --- 生成会议纪要 API (直接使用 LLM 模型) ---
@app.get("/generate_minutes")
async def generate_minutes_endpoint(meeting_id: str = Query(..., description="要生成纪要的会议 ID")):
    """
    为特定会议生成会议纪要（使用 LLM 模型生成中文摘要）。
    """
    mongodb_db = getattr(app.state, 'mongodb_manager', None)
    if not mongodb_db:
        raise HTTPException(status_code=500, detail="MongoDB 数据库未初始化。")
    if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="后端服务未初始化，请稍后重试。")

    # 确保 LLM 模型已加载
    llm_model = app.state.backend_coordinator.llm_model
    if not llm_model.is_model_loaded():
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="LLM 模型未准备就绪。")

    try:
        # 从 MongoDB 中获取转录文本
        full_transcript_entries = await mongodb_db.get_all_transcripts_for_meeting(meeting_id)

        full_transcript_text_chinese = ""
        for entry in full_transcript_entries:
            speaker = entry.speaker_id
            text = entry.text
            if isinstance(entry.timestamp, str):
               
                timestamp_dt = datetime.fromisoformat(entry.timestamp)
            else:
                timestamp_dt = entry.timestamp
            timestamp_str = timestamp_dt.strftime('%H:%M:%S')
            full_transcript_text_chinese += f"[{timestamp_str}] {speaker}: {text}\n"

        if not full_transcript_text_chinese.strip():
            raise HTTPException(status_code=404, detail="未找到该会议 ID 的转录内容。")

        # 直接使用 LLM 模型生成中文会议纪要
        logger.info(f"正在使用 LLM 模型为会议 {meeting_id} 生成简洁的会议纪要...")
        prompt = f"""
        请根据以下会议转录文本，生成一份简洁、准确的会议纪要。
        纪要应直接总结会议的核心内容，包括主要议题和重要结论。
        请使用中文，并确保信息精炼且高度相关。

        会议转录文本:
        {full_transcript_text_chinese}
        """

        minutes_content_chunks = []
        # 使用 async for 消费 generate_text 的输出
        async for chunk in llm_model.generate_text(prompt, stream=False):
            minutes_content_chunks.append(chunk)
        minutes_content_chinese = "".join(minutes_content_chunks)  # 合并所有分块

        logger.info(f"中文会议纪要生成完成（LLM）。")

        # 生成 Docx 文件
        document = docx.Document()
        document.add_heading('会议纪要', level=1)
        # 将生成的中文纪要内容作为段落添加
        # 尝试将 Markdown 的换行符转换为 Word 的段落，以获得更好的格式
        for paragraph in minutes_content_chinese.split('\n'):
            if paragraph.strip():  # 避免添加空段落
                document.add_paragraph(paragraph.strip())

        doc_io = BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        filename = f"会议纪要-{meeting_id}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"生成会议纪要失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成会议纪要失败: {e}")


# --- 导出用户发言 API ---
@app.get("/export_user_speech")
async def export_user_speech_endpoint(meeting_id: str = Query(..., description="要导出发言内容的会议 ID"), user_id: Optional[str] = Query(None, description="可选: 按用户 ID 过滤")):
    """
    从指定会议中导出用户的发言内容。
    """
    mongodb_db = getattr(app.state, 'mongodb_manager', None)
    if not mongodb_db:
        raise HTTPException(status_code=500, detail="MongoDB 数据库未初始化。")

    try:
        speeches_data = await mongodb_db.get_user_speeches_for_meeting(meeting_id, user_id)

        user_speeches: Dict[str, List[str]] = {}
        for speech_row in speeches_data:
            current_user_id = speech_row.get('user_id', 'unknown_user')
            speaker_name = speech_row.get('speaker_id', '未知发言人')
            text = speech_row.get('text', '')

            timestamp_raw = speech_row.get('timestamp')
            if isinstance(timestamp_raw, str):
                try:
                    timestamp_dt = datetime.fromisoformat(timestamp_raw)
                except ValueError:
                    logger.warning(f"时间戳 '{timestamp_raw}' 格式无效，无法解析为 datetime。使用原始字符串。")
                    timestamp_dt = None
            elif isinstance(timestamp_raw, datetime):
                timestamp_dt = timestamp_raw
            else:
                timestamp_dt = None

            timestamp_str = "N/A"
            if timestamp_dt:
                timestamp_str = timestamp_dt.strftime('%H:%M:%S')

            if current_user_id not in user_speeches:
                user_speeches[current_user_id] = []
            user_speeches[current_user_id].append(f"[{timestamp_str}] {speaker_name}: {text}")

        if not user_speeches:
            raise HTTPException(status_code=404, detail="未找到该会议 ID 或该用户的发言内容。")

        output_content = ""
        for uid, speeches in user_speeches.items():
            username_display = uid
            try:
                profile_data = await mongodb_db.get_user(uid)
                if profile_data:
                    username_display = profile_data.username
            except Exception as e:
                logger.warning(f"无法为用户 {uid} 获取用户名 (MongoDB): {e}")

            output_content += f"--- 用户: {username_display} (ID: {uid}) ---\n"
            output_content += "\n".join(speeches)
            output_content += "\n\n"

        filename = f"用户发言-{meeting_id}.txt"
        if user_id:
            filename = f"用户发言-{username_display}-{meeting_id}.txt"

        encoded_filename = quote(filename)
        return StreamingResponse(
            BytesIO(output_content.encode('utf-8')),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"导出用户发言失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"导出用户发言失败: {e}")

# --- 生成会议摘要报告 API (继续使用 LLM 模型) ---
@app.get("/generate_meeting_report")
async def generate_meeting_report_endpoint(meeting_id: str = Query(..., description="要生成报告的会议 ID")):
    """
    生成会议摘要报告（使用 LLM 模型进行结构化报告生成）。
    """
    mongodb_db = getattr(app.state, 'mongodb_manager', None)
    if not mongodb_db:
        raise HTTPException(status_code=500, detail="MongoDB 数据库未初始化。")
    if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="后端服务未初始化，请稍后重试。")

    try:
        # 从 MongoDB 中获取转录文本
        full_transcript_entries = await mongodb_db.get_all_transcripts_for_meeting(meeting_id)

        full_transcript_text = ""
        for entry in full_transcript_entries:
            speaker = entry.speaker_id
            text = entry.text
            if isinstance(entry.timestamp, str):
               
                timestamp_dt = datetime.fromisoformat(entry.timestamp)
            else:
                timestamp_dt = entry.timestamp
            timestamp_str = timestamp_dt.strftime('%H:%M:%S')
            full_transcript_text += f"[{timestamp_str}] {speaker}: {text}\n"

        if not full_transcript_text.strip():
            raise HTTPException(status_code=404, detail="未找到该会议 ID 的转录内容。")

        # 继续使用 LLM 模型，因为它能处理更复杂的结构化报告生成和信息提取，而不仅仅是纯粹的总结。
        # 优化 LLM 提示，使其更关注准确性和结构化输出
        prompt = f"""
        请根据以下会议转录文本，生成一份详细、准确且结构化的会议摘要报告。
        报告应包含以下部分：
        1.  **会议主题**: 总结会议的核心主题。
        2.  **关键讨论点**: 列出会议中讨论的所有重要问题。
        3.  **关键决策**: 清晰地说明会议中做出的所有决定。
        4.  **行动事项**: 列出所有需要执行的任务，包括负责人员和截止日期（如果转录中提及）。
        5.  **后续步骤**: 概述会议后需要进行的跟进工作。

        请使用清晰的 Markdown 格式，并确保信息准确无误，避免遗漏关键细节。

        会议转录文本:
        {full_transcript_text}
        """
        logger.info(f"正在使用 LLM 模型为会议 {meeting_id} 生成详细摘要报告...")

        report_content_chunks = []
        # 使用 async for 消费 generate_text 的输出
        async for chunk in app.state.backend_coordinator.llm_model.generate_text(prompt, stream=False):
            report_content_chunks.append(chunk)
        report_content_markdown = "".join(report_content_chunks)  # 合并所有分块

        logger.info(f"详细摘要报告生成完成（LLM）。")

        # 以 JSON 格式返回报告内容（前端可以解析和显示或下载）
        return JSONResponse(
            content={"message": "会议摘要报告生成成功", "report": {"content": report_content_markdown}},
            media_type="application/json"
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"生成会议摘要报告失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成会议摘要报告失败: {e}")

# 触发知识图谱集成的接口
@app.post("/link_meeting_to_kg")
async def link_meeting_to_kg_endpoint(meeting_id: str = Query(..., description="要链接到知识图谱的会议 ID")):
    """
    从指定会议的转录文本中提取知识图谱信息并存储到 Neo4j 中。
    """
    if not hasattr(app.state, 'backend_coordinator') or app.state.backend_coordinator is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="后端服务未初始化，请稍后重试。")

    try:
        result = await app.state.backend_coordinator.process_meeting_for_knowledge_graph(meeting_id)
        return JSONResponse(content=result, status_code=status.HTTP_200_OK)
    except RuntimeError as e:
        logger.error(f"将会议 '{meeting_id}' 链接到知识图谱失败: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"将会议链接到知识图谱失败: {e}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"将会议 '{meeting_id}' 链接到知识图谱时发生未知错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"发生未知错误: {e}")

@app.get("/ask_question")
async def ask_question_endpoint(question: str, meeting_id: str = Query(..., description="要提问的会议 ID")):
    """
    接收用户关于会议内容的问题，并由智能会议助手提供答案。
    """
    try:
        if not hasattr(app.state, 'backend_coordinator') or \
           app.state.backend_coordinator is None or \
           not hasattr(app.state.backend_coordinator, 'meeting_assistant') or \
           app.state.backend_coordinator.meeting_assistant is None:
            raise HTTPException(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                detail="会议助手服务未初始化，请稍后重试。"
            )
        # 注意: meeting_assistant.answer_question 内部调用 LLM，但它应该返回一个字符串。
        # 如果您希望此接口也支持流式传输，则需要修改 meeting_assistant.answer_question 的返回类型
        # 并将此接口也更改为 StreamingResponse。
        answer = await app.state.backend_coordinator.meeting_assistant.answer_question(meeting_id, question)
        logger.info(f"问答请求: '{question[:50]}...'，答案: '{answer[:50]}...'")
        return {
            "answer": answer,
            "timestamp": datetime.now().isoformat()
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"会议问答失败（问题: '{question}'）: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"会议问答失败: {str(e)}"
        )

@app.websocket("/monitor-ws")
async def monitor_websocket_endpoint(websocket: WebSocket):
    """
    为监控仪表板提供 WebSocket 连接，以获取实时系统状态更新。
    """
    if not hasattr(app.state, 'monitor_manager') or app.state.monitor_manager is None:
        logger.error("monitor_manager 未初始化，无法建立监控 WebSocket 连接。")
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="服务器内部服务未准备就绪。")

    await app.state.monitor_manager.connect(websocket)

    try:
        while True:
            await app.state.monitor_manager.send_system_status(websocket)
            await asyncio.sleep(settings.MONITOR_BROADCAST_INTERVAL)

    except WebSocketDisconnect:
        logger.info("监控客户端已断开连接")
    except Exception as e:
        logger.error(f"监控连接错误: {str(e)}", exc_info=True)
        await websocket.close(code=status.WS_1011_INTERNAL_ERROR)
    finally:
        if hasattr(app.state, 'monitor_manager') and app.state.monitor_manager is not None:
            app.state.monitor_manager.disconnect(websocket)

